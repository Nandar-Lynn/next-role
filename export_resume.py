import argparse
import re
from pathlib import Path


def markdown_to_text(markdown: str) -> str:
    text = markdown
    # Remove images
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    # Convert links to their text
    text = re.sub(r"\[(.*?)\]\((?:.*?)\)", r"\1", text)
    # Remove headings
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    # Convert bold/italic markers
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    # Convert lists
    text = re.sub(r"^\s*[-*+]\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "- ", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # Remove code fences and backticks
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    # Clean repeated blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


def write_txt(path: Path, markdown: str) -> None:
    text = markdown_to_text(markdown)
    path.write_text(text, encoding="utf-8")


def write_docx(path: Path, markdown: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to export DOCX. Install it with: pip install python-docx"
        ) from exc

    text = markdown_to_text(markdown)
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def parse_args():
    parser = argparse.ArgumentParser(description="Export a Markdown resume to TXT or DOCX.")
    parser.add_argument("--input", default="final_resume.md", help="Path to the input Markdown resume.")
    parser.add_argument("--output", default="final_resume.txt", help="Output path (.txt or .docx).")
    return parser.parse_args()


def main():
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    markdown = input_path.read_text(encoding="utf-8")
    if output_path.suffix.lower() == ".txt":
        write_txt(output_path, markdown)
    elif output_path.suffix.lower() == ".docx":
        write_docx(output_path, markdown)
    else:
        raise ValueError("Output file must use .txt or .docx extension.")

    print(f"Exported {input_path} to {output_path}")


if __name__ == "__main__":
    main()
